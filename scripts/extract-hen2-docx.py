#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image


CHOICE_RE = re.compile(r"(?:^|[\s|]+)([A-Fa-f])\s*[.)]\s*")
ANSWER_RE = re.compile(r"^\s*([A-Fa-f])\s*[.)]")
QUESTION_NO_RE = re.compile(r"^\s*(\d+)\s*\.?\s*$")


def clean_text(text: str, *, preserve_newlines: bool = True) -> str:
  text = text.replace("\xa0", " ").replace("\r", "\n")
  text = re.sub(r"[ \t]+\n", "\n", text)
  text = re.sub(r"\n[ \t]+", "\n", text)
  text = re.sub(r"(?<=[a-z0-9])\.(?=[A-Z])", ". ", text)
  text = re.sub(r"(?<=[a-z0-9])\?(?=[A-Z])", "? ", text)
  if preserve_newlines:
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
  else:
    text = re.sub(r"\s+", " ", text)
  return text.strip()


def parse_question_no(text: str) -> int | None:
  match = QUESTION_NO_RE.match(text)
  return int(match.group(1)) if match else None


def parse_choices(text: str) -> list[tuple[str, str]]:
  matches = list(CHOICE_RE.finditer(text))
  choices: list[tuple[str, str]] = []
  for index, match in enumerate(matches):
    start = match.end()
    end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
    choice = clean_text(text[start:end], preserve_newlines=False)
    if choice:
      choices.append((match.group(1).upper(), choice))
  return choices


def parse_answer_letter(text: str) -> str | None:
  match = ANSWER_RE.match(text)
  return match.group(1).upper() if match else None


def clean_explanation(text: str) -> str:
  text = clean_text(text)
  lines = [line.strip() for line in text.splitlines()]
  while lines and not lines[0]:
    lines.pop(0)
  if lines and lines[0].lower().startswith("answer:"):
    lines.pop(0)
  if lines and lines[0].strip().lower() == "explanation":
    lines.pop(0)
  return "\n".join(line for line in lines if line).strip()


def image_parts_for_row(doc, row) -> Iterable[object]:
  rels = doc.part.rels
  for cell_index, cell in enumerate(row.cells):
    if cell_index != 1:
      continue
    for blip in cell._tc.xpath(".//a:blip"):
      rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
      if not rel_id:
        continue
      rel = rels[rel_id]
      yield getattr(rel, "target_part", None) or rel._target


def write_question_image(part, question_no: int, seq: int, public_dir: Path, public_url: str) -> str:
  suffix = Path(str(part.partname)).suffix.lower() or ".png"
  filename = f"q{question_no:03d}-{seq}{suffix}"
  out_path = public_dir / filename
  out_path.write_bytes(part.blob)
  with Image.open(out_path) as image:
    width, height = image.size
  src = f"{public_url.rstrip('/')}/{filename}"
  alt = f"Question {question_no} figure {seq}"
  return f"[[image:{src}|{width}x{height}|{alt}]]"


def build_workbook(rows: list[list[str | int | None]], out_xlsx: Path) -> None:
  wb = Workbook()
  ws = wb.active
  ws.title = "HEN-2 Mock"

  header = [
    "question",
    "A",
    "B",
    "C",
    "D",
    "E",
    "F",
    "correct",
    "explanation",
    "topic",
    "difficulty",
  ]
  ws.append(header)
  for row in rows:
    ws.append(row)

  header_fill = PatternFill("solid", fgColor="1F2937")
  for cell in ws[1]:
    cell.font = Font(color="FFFFFF", bold=True)
    cell.fill = header_fill
    cell.alignment = Alignment(vertical="center", wrap_text=True)

  widths = {
    "A": 74,
    "B": 30,
    "C": 30,
    "D": 30,
    "E": 30,
    "F": 30,
    "G": 30,
    "H": 10,
    "I": 78,
    "J": 42,
    "K": 12,
  }
  for column, width in widths.items():
    ws.column_dimensions[column].width = width
  for row in ws.iter_rows(min_row=2):
    for cell in row:
      cell.alignment = Alignment(vertical="top", wrap_text=True)
  ws.freeze_panes = "A2"
  ws.auto_filter.ref = f"A1:{get_column_letter(ws.max_column)}{ws.max_row}"
  wb.save(out_xlsx)


def extract(args: argparse.Namespace) -> None:
  question_doc = Document(args.questions_docx)
  answer_doc = Document(args.answer_key_docx)
  question_table = question_doc.tables[0]
  answer_table = answer_doc.tables[0]

  args.public_dir.mkdir(parents=True, exist_ok=True)

  rows: list[list[str | int | None]] = []
  topic = ""
  image_count = 0
  warnings: list[str] = []

  for row_index, (question_row, answer_row) in enumerate(
    zip(question_table.rows[1:], answer_table.rows[1:]),
    start=1,
  ):
    question_no = parse_question_no(question_row.cells[0].text)
    if question_no is None:
      topic = clean_text(question_row.cells[0].text, preserve_newlines=False)
      continue

    answer_no = parse_question_no(answer_row.cells[0].text)
    if answer_no != question_no:
      raise ValueError(f"row {row_index}: question/answer number mismatch {question_no} != {answer_no}")

    stem = clean_text(question_row.cells[1].text)
    image_tokens: list[str] = []
    for seq, part in enumerate(image_parts_for_row(question_doc, question_row), start=1):
      image_tokens.append(write_question_image(part, question_no, seq, args.public_dir, args.public_url))
    if image_tokens:
      image_count += len(image_tokens)
      stem = f"{stem}\n\n" + "\n\n".join(image_tokens)

    choices = parse_choices(question_row.cells[2].text)
    if len(choices) < 2:
      warnings.append(f"Q{question_no}: only {len(choices)} choices parsed")
      continue

    correct_letter = parse_answer_letter(answer_row.cells[2].text)
    choice_letters = [letter for letter, _ in choices]
    if correct_letter not in choice_letters:
      warnings.append(
        f"Q{question_no}: answer {correct_letter or '?'} not present in parsed choices {choice_letters}",
      )
      continue

    explanation = clean_explanation(answer_row.cells[3].text)
    row: list[str | int | None] = [
      stem,
      None,
      None,
      None,
      None,
      None,
      None,
      correct_letter,
      explanation,
      topic,
      None,
    ]
    for column_offset, (_letter, choice) in enumerate(choices[:6], start=1):
      row[column_offset] = choice
    rows.append(row)

  if len(rows) != 151:
    warnings.append(f"expected 151 questions, wrote {len(rows)}")

  build_workbook(rows, args.out_xlsx)

  print(f"wrote {args.out_xlsx}")
  print(f"copied {image_count} image(s) to {args.public_dir}")
  print(f"extracted {len(rows)} question(s)")
  for warning in warnings:
    print(f"warning: {warning}")


def main() -> None:
  parser = argparse.ArgumentParser(description="Extract HEN-2 DOCX tables into WilliamsPod XLSX format.")
  parser.add_argument("--questions-docx", type=Path, required=True)
  parser.add_argument("--answer-key-docx", type=Path, required=True)
  parser.add_argument("--out-xlsx", type=Path, default=Path("hen-2-mock.xlsx"))
  parser.add_argument("--public-dir", type=Path, default=Path("public/hen-2"))
  parser.add_argument("--public-url", default="/hen-2")
  extract(parser.parse_args())


if __name__ == "__main__":
  main()
